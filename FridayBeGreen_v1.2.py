#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Author : N0Coriander
# address : https://github.com/N0Coriander
# Date : 2021/6/18 13:46
# Desc : 将分析报告粘贴到脚本同级目录下，运行程序后，会遍历所有word文档，自动形成周报数据所需格式，然后塞到剪贴板中，支持多文件读取与数据处理，满足周报填写格式
# Compare : 根据项目名称获取其所属地区(省/市)及客户行业
"""
【所需参数】
0 项目 取客户名称中第一个下划线之前的
1 地区 本地建txt，读取里面的内容，如果是新项目让用户自己手动输入行业，然后我追加到txt中
2 行业 本地建txt，读取里面的内容，如果是新项目让用户自己手动输入行业，然后我追加到txt中
3 服务类型 定期分析或售前，根据02和01区分
4 服务周期，如果是定期，就是每月，如果是售前，就为空
5 安全分析工程师 取表格里的名字
6 审核人 取表格里的名字

7 事件报告 0
8 定期报告 1
9 总计 1

10 告警数量 获取第一段中的那个数
11 事件输出量 获取事件列表中的行数，试着获取1.1章节第二段开头就有这个数量
12 事件交付量 获取事件列表中的行数，试着获取1.1章节第二段开头就有这个数量
13 含人工发掘事件量 0
14 含平台误报事件 0
15 含平台漏报事件 0
16 总计 获取事件列表中的行数，试着获取1.1章节第二段开头就有这个数量

17-22 分析沟通时长 日常支持工作时长 分析时长 审核时长 修订时长 总计 均为空

23-27 本周故障次数 故障类型说明 解决时长(h) 产品BUG/需求 jira/tac链接 均为0

28-37 1名称 数量 2名称 数量 3名称 数量 4名称 数量 5名称 数量 获取插入的网络攻击excel中的前五行，如果不够五行，用-填充

38-47 1名称 数量 2名称 数量 3名称 数量 4名称 数量 5名称 数量 获取插入的恶意软件excel中的前五行，如果不够五行，用-填充
"""
# 自动下载并导入pyperclip、python-docx、pandas
import re
import zipfile
import shutil
try:
    import os
    import pandas as pd
    import pyperclip
    from docx import Document
except ImportError:
    import os
    os.system('pip3 install pandas==1.1.3')
    os.system('pip3 install pyperclip')
    os.system('pip3 install python-docx')
    import pandas as pd
    import pyperclip
    from docx import Document
input('[+] 请先将本周用于填写周报数据的分析报告粘贴在脚本所在目录下，然后按回车键开始')

files = os.listdir()  # 列出目录下所有文件
mirror = ["['']", ""]
# 48个，放单个报告的数据
end = ['', '', '', '', '', '', '',
       '0', '1', '1',
       '', '', '', '0', '0', '0', '',
       '', '', '', '', '', '',
       '0', '0', '0', '0', '0',
       '', '', '', '', '', '', '', '', '', '',
       '', '', '', '', '', '', '', '', '', '', ]
# 放所有word文档的数据
end_list = []

# 定一个空列表用于存放需要处理的word文档名称
word_file = []
for i in files:
    if '.docx' in i:
        word_file.append(i)
    else:
        pass
# 判断列表中是否存在元素为空或者是空列表（windows中运行时，会不自觉插入空列表元素）
for q in mirror:
    if q in word_file:
        word_file.remove(q)
    else:
        pass
# 判断是否复制文件进来
if len(word_file) < 1:
    print('[ERROR] 大哥你先复制文件进来呀！')
else:
    pass

MSS_TDR_Projects = ''
project_area = ''   # 客户地区
project_type = ''   # 客户类型
# print(word_file)
# 开始获取每个word文档中的所需数据，n为文档名称
for n in word_file:
    # 客户名称end[0]
    no_docx = n[0:-5]   # 去掉末尾的.docx
    n_list = no_docx.split('_')
    end[0] = n_list[0]

    # 获取客户地区end[1]、行业end[2]
    with open('MSS-TDR分析项目.txt') as mss:
        MSS_TDR_Projects = mss.read().strip()
    if end[0] in MSS_TDR_Projects:
        for txt_line in MSS_TDR_Projects.split('\n'):
            if end[0] in txt_line:
                end[1] = txt_line.split('、')[1]
                end[2] = txt_line.split('、')[2]
            else:
                pass
    else:
        print(f'[!] 检测到{end[0]}为新客户！')
        project_area = input('[+] 请输入客户所在地区：')
        project_type = input('[+] 请输入客户所属行业：')
        end[1] = project_area
        end[2] = project_type
        with open('MSS-TDR分析项目.txt', 'a') as new_mss:
            new_mss.write('\n' + end[0] + '、' + project_area + '、' + project_type)

    # 获取服务类型end[3]、服务周期end[4]
    if n_list[3][0:2] == '02':
        end[3] = '定期分析'
        end[4] = '每月'
    elif n_list[3][0:2] == '01':
        end[3] = '售前测试'
        end[4] = ''
    else:
        print('[ERROR] 大哥你文件名里面的编号有点问题呀！')

    # 获取安全分析工程师人名end[5]、审核人end[6]
    document = Document(n)
    tables = document.tables  # 获取文档中所有表格对象，放到一个列表里
    first_table_its_name_is_WendangXinxi = tables[0]  # 获取word文档中的第一个表格：文档信息
    first_table_its_name_is_WendangXinxi_cells = first_table_its_name_is_WendangXinxi._cells    # 用_cells方法获取文档信息表格中所有的单元格
    first_table_its_name_is_WendangXinxi_cells_info = [cell.text for cell in first_table_its_name_is_WendangXinxi_cells]    # 获取单元格内所有的文字信息
    end[5] = first_table_its_name_is_WendangXinxi_cells_info[13]
    end[6] = first_table_its_name_is_WendangXinxi_cells_info[17]

    # 获取告警数量end[10]，位置是在1简报下的1.1安全事件总结第一段中的第一行
    paragraphs = document.paragraphs    # 获取文档中的所有段落
    all_paragraphs_info = [par.text for par in paragraphs]  # 获取所有段落文字信息
    an_quan_shi_jian_zongjie_the_first_paragraph_index = all_paragraphs_info.index('安全事件总结') + 1   # 获取1.1章节中第一段落的索引号
    an_quan_shi_jian_zongjie_the_first_paragraph = all_paragraphs_info[an_quan_shi_jian_zongjie_the_first_paragraph_index]    # 获取1.1章节中第一段落的文字内容
    the_first_paragraph_find_nmuber_expression = re.compile(r'(\d+)')   # 筛选告警数量
    end[10] = the_first_paragraph_find_nmuber_expression.findall(an_quan_shi_jian_zongjie_the_first_paragraph)[0]

    # 事件输出量、事件交付量、总结分别是end[11]、end[12]、end[16]，位置在1.1章节第二段开头就有这个数量
    an_quan_shi_jian_zongjie_the_second_paragraph_index = all_paragraphs_info.index('安全事件总结') + 2  # 获取1.1章节中第二段落的索引号
    an_quan_shi_jian_zongjie_the_second_paragraph = all_paragraphs_info[an_quan_shi_jian_zongjie_the_second_paragraph_index]  # 获取1.1章节中第二段落的文字内容
    the_second_paragraph_find_nmuber_expression = re.compile(r'(\d+)')  # 筛选告警数量
    # 适配报告中没有事件时的场景
    if len(the_second_paragraph_find_nmuber_expression.findall(an_quan_shi_jian_zongjie_the_second_paragraph)) > 0:
        end[11] = the_second_paragraph_find_nmuber_expression.findall(an_quan_shi_jian_zongjie_the_second_paragraph)[0]
        end[12] = the_second_paragraph_find_nmuber_expression.findall(an_quan_shi_jian_zongjie_the_second_paragraph)[0]
        end[16] = the_second_paragraph_find_nmuber_expression.findall(an_quan_shi_jian_zongjie_the_second_paragraph)[0]
    else:
        end[11] = '0'
        end[12] = '0'
        end[16] = '0'

    # 28-37 1名称 数量 2名称 数量 3名称 数量 4名称 数量 5名称 数量 获取插入的网络攻击excel中的前五行，如果不够五行，用-填充
    # 38-47 1名称 数量 2名称 数量 3名称 数量 4名称 数量 5名称 数量 获取插入的恶意软件excel中的前五行，如果不够五行，用-填充
    # 1、有网络攻击事件，有恶意软件事件；2、有网络攻击事件，无恶意软件事件；3、无网络攻击事件，有恶意软件事件；4、二者均无。
    # 第四种情况：二者均无
    if '网络攻击事件' not in all_paragraphs_info:
        if '恶意软件事件' not in all_paragraphs_info:
            end[28] = end[29] = end[30] = end[31] = end[32] = end[33] = end[34] = end[35] = end[36] = end[37] = end[38] = end[39] = end[40] = end[41] = end[42] = end[43] = end[44] = end[45] = end[46] = end[47] = '-'
        else:
            pass
    else:
        pass

    # 固定的五个表Microsoft_Excel____.xlsx到Microsoft_Excel____4.xlsx，路径是/word/embeddings/，如果你在word中自己插了excel，那就会改变后面的顺序
    # 网络攻击事件和恶意软件事件表中有固定字符串：网络攻击类型分布、恶意软件类型分布
    word = zipfile.ZipFile(n)   # 因为word文档本身就是一个xml集合的压缩文件，这里进行读取
    word_xml_names = word.namelist()    # 获取到所有的文件名（包含了路径）
    Microsoft_Excel_list = []
    for word_xml_name in word_xml_names:
        if 'word/embeddings/Microsoft_Excel' in word_xml_name:
            for num in range(5, 100):
                if str(num) in word_xml_name:
                    Microsoft_Excel_list.append(word_xml_name)
                else:
                    pass
        else:
            pass
    # 判断列表中是否存在元素为空或者是空列表（windows中运行时，会不自觉插入空列表元素）
    for q in mirror:
        if q in Microsoft_Excel_list:
            Microsoft_Excel_list.remove(q)
        else:
            pass
    # print(Microsoft_Excel_list)
    if os.path.exists('temp'):
        shutil.rmtree('temp')
    else:
        pass
    os.mkdir('temp')    # 创建temp临时文件夹用于存放我要获取数据的excel表
    for Microsoft_Excel in Microsoft_Excel_list:
        word.extract(Microsoft_Excel, 'temp')   # 将文件解压缩到temp文件夹下
    word.close()    # 关闭压缩
    for p in Microsoft_Excel_list:
        Microsoft_Excel_path = 'temp/' + p
        # 1、有网络攻击事件，有恶意软件事件；2、有网络攻击事件，无恶意软件事件；3、无网络攻击事件，有恶意软件事件。
        # 网络攻击类型分布、恶意软件类型分布
        if '恶意软件类型分布' in pd.read_excel(Microsoft_Excel_path):
            # 判断够不够五行
            if len(pd.read_excel(Microsoft_Excel_path)) > 4:
                end[38] = pd.read_excel(Microsoft_Excel_path).loc[0][0].replace('【恶意软件】', '')
                end[39] = pd.read_excel(Microsoft_Excel_path).loc[0][1]
                end[40] = pd.read_excel(Microsoft_Excel_path).loc[1][0].replace('【恶意软件】', '')
                end[41] = pd.read_excel(Microsoft_Excel_path).loc[1][1]
                end[42] = pd.read_excel(Microsoft_Excel_path).loc[2][0].replace('【恶意软件】', '')
                end[43] = pd.read_excel(Microsoft_Excel_path).loc[2][1]
                end[44] = pd.read_excel(Microsoft_Excel_path).loc[3][0].replace('【恶意软件】', '')
                end[45] = pd.read_excel(Microsoft_Excel_path).loc[3][1]
                end[46] = pd.read_excel(Microsoft_Excel_path).loc[4][0].replace('【恶意软件】', '')
                end[47] = pd.read_excel(Microsoft_Excel_path).loc[4][1]
                break
            elif len(pd.read_excel(Microsoft_Excel_path)) == 4:
                end[38] = pd.read_excel(Microsoft_Excel_path).loc[0][0].replace('【恶意软件】', '')
                end[39] = pd.read_excel(Microsoft_Excel_path).loc[0][1]
                end[40] = pd.read_excel(Microsoft_Excel_path).loc[1][0].replace('【恶意软件】', '')
                end[41] = pd.read_excel(Microsoft_Excel_path).loc[1][1]
                end[42] = pd.read_excel(Microsoft_Excel_path).loc[2][0].replace('【恶意软件】', '')
                end[43] = pd.read_excel(Microsoft_Excel_path).loc[2][1]
                end[44] = pd.read_excel(Microsoft_Excel_path).loc[3][0].replace('【恶意软件】', '')
                end[45] = pd.read_excel(Microsoft_Excel_path).loc[3][1]
                end[46] = end[47] = '-'
                break
            elif len(pd.read_excel(Microsoft_Excel_path)) == 3:
                end[38] = pd.read_excel(Microsoft_Excel_path).loc[0][0].replace('【恶意软件】', '')
                end[39] = pd.read_excel(Microsoft_Excel_path).loc[0][1]
                end[40] = pd.read_excel(Microsoft_Excel_path).loc[1][0].replace('【恶意软件】', '')
                end[41] = pd.read_excel(Microsoft_Excel_path).loc[1][1]
                end[42] = pd.read_excel(Microsoft_Excel_path).loc[2][0].replace('【恶意软件】', '')
                end[43] = pd.read_excel(Microsoft_Excel_path).loc[2][1]
                end[44] = end[45] = end[46] = end[47] = '-'
                break
            elif len(pd.read_excel(Microsoft_Excel_path)) == 2:
                end[38] = pd.read_excel(Microsoft_Excel_path).loc[0][0].replace('【恶意软件】', '')
                end[39] = pd.read_excel(Microsoft_Excel_path).loc[0][1]
                end[40] = pd.read_excel(Microsoft_Excel_path).loc[1][0].replace('【恶意软件】', '')
                end[41] = pd.read_excel(Microsoft_Excel_path).loc[1][1]
                end[42] = end[43] = end[44] = end[45] = end[46] = end[47] = '-'
                break
            elif len(pd.read_excel(Microsoft_Excel_path)) == 1:
                end[38] = pd.read_excel(Microsoft_Excel_path).loc[0][0].replace('【恶意软件】', '')
                end[39] = pd.read_excel(Microsoft_Excel_path).loc[0][1]
                end[40] = end[41] = end[42] = end[43] = end[44] = end[45] = end[46] = end[47] = '-'
                break
            else:
                print('[ERROR] 大哥你word文档中插入的恶意软件类型excel文件有问题！')
        else:
            end[38] = end[39] = end[40] = end[41] = end[42] = end[43] = end[44] = end[45] = end[46] = end[47] = '-'
    for ph in Microsoft_Excel_list:
        Microsoft_Excel_path_2 = 'temp/' + ph
        # 1、有网络攻击事件，有恶意软件事件；2、有网络攻击事件，无恶意软件事件；3、无网络攻击事件，有恶意软件事件。
        # 网络攻击类型分布、恶意软件类型分布
        if '网络攻击类型分布' in pd.read_excel(Microsoft_Excel_path_2):
            # 判断够不够五行
            if len(pd.read_excel(Microsoft_Excel_path_2)) > 4:
                end[28] = pd.read_excel(Microsoft_Excel_path_2).loc[0][0].replace('【攻击利用】', '')
                end[29] = pd.read_excel(Microsoft_Excel_path_2).loc[0][1]
                end[30] = pd.read_excel(Microsoft_Excel_path_2).loc[1][0].replace('【攻击利用】', '')
                end[31] = pd.read_excel(Microsoft_Excel_path_2).loc[1][1]
                end[32] = pd.read_excel(Microsoft_Excel_path_2).loc[2][0].replace('【攻击利用】', '')
                end[33] = pd.read_excel(Microsoft_Excel_path_2).loc[2][1]
                end[34] = pd.read_excel(Microsoft_Excel_path_2).loc[3][0].replace('【攻击利用】', '')
                end[35] = pd.read_excel(Microsoft_Excel_path_2).loc[3][1]
                end[36] = pd.read_excel(Microsoft_Excel_path_2).loc[4][0].replace('【攻击利用】', '')
                end[37] = pd.read_excel(Microsoft_Excel_path_2).loc[4][1]
                break
            elif len(pd.read_excel(Microsoft_Excel_path_2)) == 4:
                end[28] = pd.read_excel(Microsoft_Excel_path_2).loc[0][0].replace('【攻击利用】', '')
                end[29] = pd.read_excel(Microsoft_Excel_path_2).loc[0][1]
                end[30] = pd.read_excel(Microsoft_Excel_path_2).loc[1][0].replace('【攻击利用】', '')
                end[31] = pd.read_excel(Microsoft_Excel_path_2).loc[1][1]
                end[32] = pd.read_excel(Microsoft_Excel_path_2).loc[2][0].replace('【攻击利用】', '')
                end[33] = pd.read_excel(Microsoft_Excel_path_2).loc[2][1]
                end[34] = pd.read_excel(Microsoft_Excel_path_2).loc[3][0].replace('【攻击利用】', '')
                end[35] = pd.read_excel(Microsoft_Excel_path_2).loc[3][1]
                end[36] = end[37] = '-'
                break
            elif len(pd.read_excel(Microsoft_Excel_path_2)) == 3:
                end[28] = pd.read_excel(Microsoft_Excel_path_2).loc[0][0].replace('【攻击利用】', '')
                end[29] = pd.read_excel(Microsoft_Excel_path_2).loc[0][1]
                end[30] = pd.read_excel(Microsoft_Excel_path_2).loc[1][0].replace('【攻击利用】', '')
                end[31] = pd.read_excel(Microsoft_Excel_path_2).loc[1][1]
                end[32] = pd.read_excel(Microsoft_Excel_path_2).loc[2][0].replace('【攻击利用】', '')
                end[33] = pd.read_excel(Microsoft_Excel_path_2).loc[2][1]
                end[34] = end[35] = end[36] = end[37] = '-'
                break
            elif len(pd.read_excel(Microsoft_Excel_path_2)) == 2:
                end[28] = pd.read_excel(Microsoft_Excel_path_2).loc[0][0].replace('【攻击利用】', '')
                end[29] = pd.read_excel(Microsoft_Excel_path_2).loc[0][1]
                end[30] = pd.read_excel(Microsoft_Excel_path_2).loc[1][0].replace('【攻击利用】', '')
                end[31] = pd.read_excel(Microsoft_Excel_path_2).loc[1][1]
                end[32] = end[33] = end[34] = end[35] = end[36] = end[37] = '-'
                break
            elif len(pd.read_excel(Microsoft_Excel_path_2)) == 1:
                end[28] = pd.read_excel(Microsoft_Excel_path_2).loc[0][0].replace('【攻击利用】', '')
                end[29] = pd.read_excel(Microsoft_Excel_path_2).loc[0][1]
                end[30] = end[31] = end[32] = end[33] = end[34] = end[35] = end[36] = end[37] = '-'
                break
            else:
                print('[ERROR] 大哥你word文档中插入的网络攻击类型excel文件有问题！')
        else:
            end[28] = end[29] = end[30] = end[31] = end[32] = end[33] = end[34] = end[35] = end[36] = end[37] = '-'
    shutil.rmtree('temp', ignore_errors=True)
    # （1）用'	'做连接符，复制到excel中时，便于每个值放入到每个单元格中；（2）然后插入至end_list，用于输出多文档时的数据
    jian_tie_ban = '	'.join(str(i) for i in end)
    end_list.append(jian_tie_ban)

# 判断列表中是否存在元素为空或者是空列表（windows中运行时，会不自觉插入空列表元素）
for q in mirror:
    if q in end_list:
        end_list.remove(q)
    else:
        pass

# 换行分割放入剪贴板中
pyperclip.copy('\n'.join(end_list))
print('[!] 内容已提取到剪贴板，请自主粘贴')

